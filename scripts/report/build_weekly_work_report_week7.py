from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REFERENCE = Path("reports/服饰实例分割项目周报-第六周.docx")
OUTPUT = Path("reports/服饰实例分割项目周报-第七周.docx")
FONT_NAME = "Hiragino Sans GB"
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(92, 92, 92)


def set_run_font(run, size: float, *, bold: bool = False, color=None) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(round(width_inches * 1440)))


def set_cell_margins(cell, *, top: int = 55, start: int = 70, bottom: int = 55, end: int = 70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_table(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tbl_pr.append(table_width)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 8.8,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold)


def add_paragraph(doc: Document, text: str, *, size: float = 10.8):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2.5)
    paragraph.paragraph_format.line_spacing = 1.04
    run = paragraph.add_run(text)
    set_run_font(run, size)
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = 1.02
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.8)
    normal.paragraph_format.line_spacing = 1.04
    normal.paragraph_format.space_after = Pt(2.5)
    for style_name, size in (("Heading 1", 14.2), ("Heading 2", 12.5)):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.font.bold = True
        style.paragraph_format.space_before = Pt(4.5)
        style.paragraph_format.space_after = Pt(1.5)
        style.paragraph_format.keep_with_next = True


def clear_template_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1.5)
    run = title.add_run("本周工作汇报：服饰细粒度视觉模块（3.1.2 选择器收敛验证）")
    set_run_font(run, 16, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run("汇报周期：第七周｜本周目标：独立 OOF Hit@0.3 达到 60%")
    set_run_font(run, 10.2, color=GRAY)


def add_protocol_table(doc: Document) -> None:
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    widths = [1.0, 2.35, 1.0, 2.35]
    values = [
        ("评估数据", "2,338 条弱标签记录，1,529 张图", "验证方式", "按图像分组的 5-fold OOF"),
        ("防泄漏", "候选与特征不使用 target bbox", "人工基准", "161 条审校框冻结，不参与调参"),
    ]
    for row_cells, row_values in zip(table.rows, values, strict=True):
        for index, (cell, value) in enumerate(zip(row_cells.cells, row_values, strict=True)):
            set_cell_text(cell, value, bold=index % 2 == 0, size=8.6)
            if index % 2 == 0:
                set_cell_shading(cell, "F2F4F7")
    configure_table(table, widths)


def add_progress_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    widths = [2.2, 0.8, 0.9, 0.9, 0.9, 1.0]
    headers = ["方案", "Hit@0.3", "命中数", "cuff", "waist", "结论"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_text(cell, header, bold=True, size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "E9EEF5")
    rows = [
        ("在线基线", "36.23%", "847", "31.16%", "65.79%", "起点"),
        ("区域条件 DINOv2 listwise", "47.78%", "1,117", "43.39%", "73.39%", "+270"),
        ("+ cuff patch 空间特征", "50.21%", "1,174", "46.24%", "73.39%", "+57"),
        ("+ 在线服饰实例几何", "51.41%", "1,202", "47.14%", "76.32%", "最佳"),
        ("+ 候选共识/专家交互 v5", "51.41%", "1,202", "47.09%", "76.61%", "无净增益"),
        ("+ cuff 成对重排", "51.33%", "1,200", "46.99%", "76.61%", "下降 2"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row, strict=True)):
            set_cell_text(
                cell,
                value,
                bold=value == "最佳",
                size=8.1,
                align=WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )
    configure_table(table, widths)


def add_cuff_diagnostic_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [2.35, 1.05, 3.3]
    headers = ["诊断项", "数量", "判断"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_text(cell, header, bold=True, size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "E9EEF5")
    rows = [
        ("cuff 当前选择命中", "940 / 1,996", "主要瓶颈仍是候选排序"),
        ("完整候选 oracle 命中", "1,286", "候选池仍有较大理论空间"),
        ("同侧候选 oracle 命中", "1,182", "强制同侧会损失 104 个 oracle 命中"),
        ("错侧失败且可由同侧恢复", "9", "硬左右过滤的可恢复收益很小"),
        ("当前错侧但已命中", "19", "硬过滤反而可能破坏已有命中"),
        ("完整左右配对 / 框碰撞", "809 / 52", "仅靠碰撞规则不足以稳定提升"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row, strict=True)):
            set_cell_text(cell, value, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER if index == 1 else WD_ALIGN_PARAGRAPH.LEFT)
    configure_table(table, widths)


def add_gap_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1.55, 1.45, 1.45, 2.25]
    headers = ["指标", "Hit@0.3", "命中数", "说明"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_text(cell, header, bold=True, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "E9EEF5")
    rows = [
        ("本周最佳 OOF", "51.41%", "1,202", "可信的当前上限"),
        ("周目标", "60.00%", "1,403", "仍差 201 个命中"),
        ("候选 oracle", "68.22%", "1,595", "理论空间存在，但当前监督难以学到"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row, strict=True)):
            set_cell_text(cell, value, size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER if index in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT)
    configure_table(table, widths)


def build_report() -> None:
    doc = Document(REFERENCE)
    clear_template_body(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("一、本周目标与评估协议", level=1)
    add_paragraph(
        doc,
        "本周目标是把 3.1.2 的 Hit@0.3 从 36.23% 提升到 60%。为回应前期关于伪标签噪声和数据泄漏的反馈，"
        "本轮固定 2,338 条 cuff/waist 弱标签记录，所有候选均由在线模型从图像生成；target bbox 只提供训练监督和最终计分，"
        "不进入候选生成或特征。人工 bbox benchmark 保持冻结，仅用于后续验真。",
    )
    add_protocol_table(doc)

    doc.add_heading("二、选择器特征增强与 OOF 结果", level=1)
    add_paragraph(
        doc,
        "在相同的 image-grouped 5-fold OOF 协议下，依次加入区域条件 DINOv2、cuff patch 空间特征和 3.1.1 在线预测服饰实例几何。"
        "其中在线服饰几何把命中数从 1,174 提升到 1,202，是本周最后一个可复现的有效增益。",
    )
    add_progress_table(doc)

    doc.add_heading("三、完成的工程与消融验证", level=1)
    add_bullet(doc, "加入 cuff 专用 DINOv2 patch 空间描述，补充候选内部的局部视觉结构。")
    add_bullet(doc, "接入 3.1.1 在线预测服饰框构造相对几何；2,338 条记录全部成功，且未使用弱标签框。")
    add_bullet(doc, "候选共识无净增益，pair reranker 少 2 个命中；164 tests passed，CUDA、固定 seed 与 OOF 可复现。")

    diagnostic_heading = doc.add_heading("四、左右袖口诊断", level=1)
    diagnostic_heading.paragraph_format.page_break_before = True
    add_paragraph(
        doc,
        "对 1,996 条 cuff 记录进一步检查左右语义。多数预测已经满足同侧约束，真正因错侧而失败且可恢复的样本只有 9 条；"
        "同时有 19 条错侧候选已经达到 Hit@0.3。因此直接加入硬左右过滤，净收益风险大于收益。",
    )
    add_cuff_diagnostic_table(doc)

    doc.add_heading("五、距离 60% 的差距与停止判断", level=1)
    add_gap_table(doc)
    add_paragraph(
        doc,
        "当前候选池的 oracle 高于 60%，说明理论上仍有空间；但连续三轮新增特征的结果已从 +28 个命中变为 0、再变为 -2。"
        "这表明瓶颈不是继续堆叠几何规则，而是带噪 landmark pseudo-label 无法稳定教会选择器识别真实局部区域。"
        "若继续以这批伪标签做目标，很可能只提高伪标签 IoU，无法证明真实定位能力提升。",
    )
    add_paragraph(
        doc,
        "因此，本周不把 60% 写成已完成目标，也不通过改划分、看测试集调阈值或引入 target bbox 特征来追数。"
        "当前可信结果冻结为 51.41%（1,202/2,338），比 60% 仍少 201 个命中。",
    )

    doc.add_heading("六、当前结论与后续计划", level=1)
    add_bullet(doc, "冻结“DINOv2 cuff 空间特征 + 在线服饰实例几何”的线性 listwise 版本，作为当前弱标签 OOF 最佳模型。")
    add_bullet(doc, "停止继续在同一批 landmark pseudo-label 上做选择器特征工程；candidate consensus 和 cuff pair reranker 记录为负向消融。")
    add_bullet(doc, "人工审校 benchmark 继续作为独立验真集，不参与阈值或模型选择，避免再次出现评估泄漏。")
    add_bullet(doc, "若后续仍要求达到 60%，需要进入新阶段：补充干净人工框、微调 dense grounding，或重构候选生成，而不是继续微调当前选择器。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
