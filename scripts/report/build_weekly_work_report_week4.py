from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("reports/服饰实例分割项目周报-第四周.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10)
    run.bold = bold


def add_paragraph(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    paragraph = add_paragraph(doc, text, "List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(1)

    for style_name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(1)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("本周工作汇报：服饰细粒度视觉模块（3.1.2）")
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(16)
    run.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("汇报周期：第四周｜方向：多模态电商服饰细粒度视觉基础模块")
    meta_run.font.name = "SimSun"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    meta_run.font.size = Pt(10.5)
    meta_run.font.color.rgb = RGBColor(90, 90, 90)


def add_baseline_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["阶段", "核心方法", "验证规模", "主要结果"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")

    rows = [
        ("开放词汇 baseline", "query parser + 候选区域 + 启发式排序", "200 图 / 600 query", "weak IoU 0.2818，Hit@0.3 0.3933"),
        ("轻量 hash ranker", "文本哈希 + 几何特征，neckline/hem 生效", "200 图 / 600 query", "weak IoU 0.2822，接近启发式"),
        ("Chinese-CLIP 零样本", "中文 CLIP 文本-区域相似度", "2000 candidate groups", "top1 IoU 0.2020，单独使用效果不足"),
        ("Listwise context ranker", "query/candidate/box/context 联合排序", "训练 50k query groups", "val top1 IoU 0.5113"),
        ("后续切片验证", "同一 checkpoint eval-only", "offset 50k / 5000 groups", "top1 IoU 0.4456，oracle 0.5193"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)


def add_pipeline_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["方案", "ranker 使用策略", "avg weak IoU", "Hit@0.3", "结论"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")

    rows = [
        ("启发式 baseline", "全部区域启发式排序", "0.2818", "0.3933", "稳定但上限有限"),
        ("Listwise 全量接入", "neckline/hem/shoulder 全部用 learned", "0.2732", "0.4150", "shoulder 与 neckline 退化"),
        ("Listwise hem-only", "hem 用 learned，其余 fallback", "0.2818", "0.4050", "当前最安全线上策略"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)


def add_data_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["数据/文件", "规模", "来源", "用途"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")

    rows = [
        ("Query records", "2,808,252 条", "DeepFashion2 train mask + landmarks", "训练 query-level 弱监督样本"),
        ("Candidate records", "5,000,000 条候选", "前 500,000 query records", "训练 listwise ranker"),
        ("弱标签来源", "2,231,694 landmark / 576,558 rule", "landmark pseudo-label 优先", "评估标签质量与覆盖率"),
        ("candidate label", "395,319 positive / 4,604,681 negative", "IoU≥0.5 为正样本", "监督候选区域排序"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)


def build_report() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("一、本周主要工作", level=1)
    add_paragraph(
        doc,
        "本周主要推进 PRD 3.1.2「语言引导的局部区域定位」。根据导师反馈，我将原先偏「固定部位分割」的思路调整为"
        "「语言引导的局部区域定位」：模型不再只支持训练时定义好的 N 个部位，而是通过自然语言 query 在服饰实例内部选择目标区域。"
        "这样可以支持方位、属性、部件和关系描述，例如左边袖口、碎花图案、拉链、口袋、外套和内搭等，扩展性明显更强。"
    )
    add_paragraph(
        doc,
        "工程上，本周完成了开放词汇 baseline、DeepFashion2 弱监督标签构建、候选区域训练数据导出、多个排序模型实验，以及 AutoDL 上的训练和评估闭环。"
        "目前 3.1.2 已经从单纯规则 baseline 进入 learned ranker 验证阶段，但 full pipeline 指标显示 learned model 需要谨慎接入。"
    )

    doc.add_heading("二、Baseline 方向调整与工程实现", level=1)
    add_bullet(doc, "明确新的 baseline 定位：输入服饰图片和自然语言 query，先复用 3.1.1 实例分割结果，再在目标服饰 mask 内生成开放词汇候选区域。")
    add_bullet(doc, "实现中文 query parser，支持 neckline、hem、shoulder、cuff、pocket、pattern、zipper、button、decoration 等区域与属性描述。")
    add_bullet(doc, "实现候选区域生成与排序接口，保留 whole garment、upper、lower、left、right、center 等通用候选，避免只依赖固定部位标签。")
    add_bullet(doc, "将本地开发与 AutoDL 训练流程打通：本地 commit 后推送，AutoDL 在 /root/projects/alibaba-ai 拉取代码并运行大规模训练。")
    add_bullet(doc, "补充 weak-label evaluation 脚本，用 DeepFashion2 validation annos 中的 mask + landmarks 对 3.1.2 输出做自动评估。")

    doc.add_heading("三、弱监督数据构建", level=1)
    add_paragraph(
        doc,
        "由于 DeepFashion2 没有直接给出「领口/下摆/肩部」局部 mask，本周继续沿用 mask + landmarks 构造 pseudo-label 的路线。"
        "其中 landmark pseudo-label 优先用于 neckline、hem、shoulder；当 landmark 信息不足时，再使用规则区域作为 fallback。"
    )
    add_data_table(doc)

    doc.add_heading("四、模型实验与阶段性结果", level=1)
    add_paragraph(
        doc,
        "本周比较了三类排序路径：启发式文本规则、轻量 learned hash ranker、以及候选级 listwise context ranker。"
        "其中 Chinese-CLIP 零样本相似度也做了验证，但单独用于候选排序时 top1 IoU 只有 0.2020，说明预训练跨模态能力不能直接替代本任务的弱监督训练。"
    )
    add_baseline_table(doc)

    add_paragraph(
        doc,
        "Listwise context ranker 在候选级验证上表现较好：50k query group 训练后，第一段 2k validation top1 IoU 达到 0.5113，"
        "后续 offset 50k 的 5000 groups 也有 0.4456，接近该候选集合 oracle 0.5193。"
        "但是接入完整 3.1.2 pipeline 后，neckline 和 shoulder 指标出现退化，因此最终采用 hem-only gate。"
    )
    add_pipeline_table(doc)

    doc.add_heading("五、问题分析", level=1)
    add_bullet(doc, "候选级 listwise ranker 有学习效果，但 full pipeline 还受 3.1.1 预测 mask、候选区域生成和 weak label 噪声共同影响。")
    add_bullet(doc, "neckline 与 shoulder 对上半身边界和 landmark 分布更敏感，候选排序分数提升不一定能转化为最终 mask IoU 提升。")
    add_bullet(doc, "hem 区域相对稳定，listwise context ranker 接入后没有明显破坏整体结果，因此暂时只对 hem 启用 learned ranker。")
    add_bullet(doc, "Chinese-CLIP 零样本排序结果低于规则 baseline，说明当前阶段更可靠的路线是弱监督训练 + 保守 fallback，而不是直接替换为通用 CLIP。")
    add_bullet(doc, "本周还修正了评估输出中的 ranker_backend 统计逻辑，使 fallback 记录真实显示为 heuristic_text_region_ranker，避免误以为所有 query 都由 learned ranker 处理。")

    doc.add_heading("六、代码与文档整理", level=1)
    add_bullet(doc, "新增训练 query records 和 candidate records 构建脚本，支持从 DeepFashion2 大规模导出弱监督训练数据。")
    add_bullet(doc, "新增 candidate baseline diagnostics，用 oracle、target-region-name、Chinese-CLIP 等结果判断候选集合上限和模型空间。")
    add_bullet(doc, "实现 candidate listwise ranker 训练、soft label 训练、context feature、eval-only 验证和 full pipeline 推理接入。")
    add_bullet(doc, "更新 README、AutoDL setup 和 3.1.2 plan，记录关键命令、实验结果和当前 hem-only online policy。")
    add_bullet(doc, "补充单元测试，覆盖 listwise checkpoint 识别、unsupported region fallback、effective backend 统计和本地推理路径。")

    doc.add_heading("七、下周计划", level=1)
    add_bullet(doc, "在 AutoDL 上重新拉取最新代码，复跑 hem-only weak evaluation，确认 ranker_backend_counts 能区分 hem learned 与 neckline/shoulder fallback。")
    add_bullet(doc, "整理 3.1.2 成功/失败可视化样例，重点分析 neckline、shoulder 为什么在 full pipeline 中低于候选级验证。")
    add_bullet(doc, "构建 100-300 条人工标注小验证集，用真实局部区域标注校准 weak-label 指标，避免只围绕 noisy pseudo-label 优化。")
    add_bullet(doc, "继续改进 pseudo-label 与候选区域生成，优先提升 shoulder/neckline，再考虑 cuff、waist 等更多部位。")
    add_bullet(doc, "评估加入图像区域 embedding 或 grounding 模型的可行性，但保持 fallback 策略，避免开放词汇查询能力退化。")
    add_paragraph(
        doc,
        "总体来看，本周完成了 3.1.2 baseline 方向的关键调整，并建立了从弱监督数据、候选训练、排序模型到 full pipeline 评估的完整闭环。"
        "当前最稳妥方案是开放词汇启发式 baseline + hem-only listwise learned ranker；下一步重点不是盲目扩大模型，而是用可视化和人工小验证集确认指标瓶颈。"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
