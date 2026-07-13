from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REFERENCE = Path("reports/服饰实例分割项目周报-第五周.docx")
OUTPUT = Path("reports/服饰实例分割项目周报-第六周.docx")
FONT_NAME = "Arial Unicode MS"
BLUE = RGBColor(46, 116, 181)


def set_run_font(run, size: float, *, bold: bool = False, color=None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(round(width_inches * 1440)))


def configure_table(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tbl_pr.append(table_width)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold)


def add_paragraph(doc: Document, text: str, *, size: float = 11.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, size)
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, 11.2)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(11.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(3)
    for style_name, size in [("Heading 1", 15), ("Heading 2", 13)]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.font.bold = True
        style.paragraph_format.space_before = Pt(5)
        style.paragraph_format.space_after = Pt(2)


def clear_template_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("本周工作汇报：服饰细粒度视觉模块（3.1.2 Gated Hybrid 验证）")
    set_run_font(run, 16, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(5)
    run = meta.add_run("汇报周期：第六周｜方向：多模态电商服饰细粒度视觉基础模块")
    set_run_font(run, 10.5, color=RGBColor(90, 90, 90))


def add_result_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [1.2, 1.15, 1.15, 1.55, 1.65]
    headers = ["方案", "评估规模", "avg bbox IoU", "Hit@0.3 / Hit@0.5", "结论"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "F2F4F7")
    rows = [
        ("heuristic-only", "171", "0.2599", "0.3918 / 0.2047", "结构区域的 control baseline"),
        ("GroundingDINO-only", "171", "0.2199", "0.2456 / 0.1813", "不能替代整条 pipeline"),
        ("fixed gated hybrid", "171", "0.3060", "0.4503 / 0.2749", "当前最佳实验策略"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            set_cell_text(cell, value)
    configure_table(table, widths)


def add_region_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1.05, 1.45, 1.45, 2.75]
    headers = ["区域", "heuristic", "GroundingDINO", "本周判断"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "F2F4F7")
    rows = [
        ("pattern", "0.3080", "0.6691", "视觉纹理定位收益明确，保持 grounding 路由"),
        ("pocket", "0.0303", "0.1024", "虽仍困难，但 grounding 优于纯几何"),
        ("zipper", "接近持平", "接近持平", "暂不接入，保持 heuristic"),
        ("结构区域", "相对稳定", "整体较弱", "其它结构区域保持 heuristic"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            set_cell_text(cell, value)
    configure_table(table, widths)


def add_confidence_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [1.15, 1.2, 1.3, 1.4, 1.65]
    headers = ["阈值", "Calibration 语义 IoU", "Grounding / fallback", "Holdout 语义 IoU", "判断"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_text(cell, header, bold=True, size=9)
        set_cell_shading(cell, "F2F4F7")
    rows = [
        ("0.0", "0.3068", "27 / 0", "0.3963", "当前固定 gate，校准集最优"),
        ("0.25", "0.2441", "19 / 8", "0.3963", "回退后 calibration 明显下降"),
        ("0.30", "0.2442", "15 / 12", "0.4035", "holdout 微升，但样本小且 calibration 变差"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            set_cell_text(cell, value, size=9)
    configure_table(table, widths)


def build_report() -> None:
    doc = Document(REFERENCE)
    clear_template_body(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("一、本周主要工作", level=1)
    add_paragraph(
        doc,
        "本周围绕 3.1.2「语言引导的局部区域定位」完成了从离线 baseline 到可复现 gated hybrid 的验证闭环。"
        "在前期人工 bbox benchmark 的基础上，继续比较 heuristic、GroundingDINO 和固定区域门控策略；同时补齐单图/批量推理、"
        "人工参考框可视化和 demo manifest 自动生成工具，使实验结果可以复现、检查和汇报。"
    )

    doc.add_heading("二、预训练 Grounding Baseline 与 Gated 策略", level=1)
    add_paragraph(
        doc,
        "在合并后的 171 条人工标注上，GroundingDINO-only 不能替代完整 pipeline，但它对视觉纹理类语义目标有明显优势。"
        "因此采用固定 gated 策略：pattern、pocket 路由到 GroundingDINO，其余结构区域仍使用 3.1.1 实例分割 + heuristic 局部定位。"
    )
    add_result_table(doc)
    add_region_table(doc)

    doc.add_heading("三、工程实现与定性验证", level=1)
    add_bullet(doc, "实现显式 gated inference/evaluator：默认 online path 不变，只有实验脚本才会调用 GroundingDINO，避免影响稳定 baseline。")
    add_bullet(doc, "新增 manifest 模式：每张图只运行存在且可见的 query，避免在无口袋图片上强行测试口袋，从而造成误导性可视化。")
    add_bullet(doc, "新增基于人工 benchmark 的 demo manifest 自动选择：按区域选择 IoU 达标的样本，输出保留 query、路由、选择 IoU 和人工参考框。")
    add_bullet(doc, "可视化中绿色 GT 为人工框、橙色为预测局部区域、蓝色为 heuristic 分支选择的服饰实例。pattern、neckline、hem、shoulder 的自动示例均完成核验。")
    add_paragraph(
        doc,
        "20 图查询测试中，60 条结构 query 走 heuristic、40 条语义 query 走 grounding，路由行为正确。"
        "Grounding 分支平均局部定位延迟约 82ms，高于 PRD 30ms 目标，因此当前仍定义为实验路径，而不是默认线上策略。"
    )

    confidence_heading = doc.add_heading("四、Confidence Fallback 验证", level=1)
    confidence_heading.paragraph_format.page_break_before = True
    add_paragraph(
        doc,
        "为避免低置信度 GroundingDINO 检测拖累结果，本周新增离线 confidence fallback 分析器。它复用已完成的 gated 和 heuristic"
        "评估 JSON，不重新运行模型；按图像划分 131 条 calibration 和 40 条 holdout，避免同一服饰图片同时参与选阈值和验阈值。"
    )
    add_confidence_table(doc)
    add_paragraph(
        doc,
        "结论：阈值 0.0 在 calibration 上最优。虽然阈值 0.30 在 14 条 holdout 语义样本上有 0.0072 的微小 IoU 上升，"
        "但其 calibration 显著下降，且样本量太小，不能作为策略更新依据。当前 GroundingDINO score 不能可靠判断何时应回退到 heuristic。"
    )

    doc.add_heading("五、当前结论", level=1)
    add_bullet(doc, "171 条人工 bbox benchmark 是当前 3.1.2 的主评估依据；pseudo-label 和 weak ranker 仅保留为历史实验与诊断工具。")
    add_bullet(doc, "当前最佳实验结果为 fixed gated hybrid：avg bbox IoU 0.3060，优于 heuristic-only 的 0.2599。")
    add_bullet(doc, "默认在线路径继续保持 heuristic-only；pattern/pocket -> GroundingDINO 只通过显式实验入口启用。")
    add_bullet(doc, "confidence fallback 没有被采纳，避免把 14 条 holdout 上的微小波动误判为真实模型提升。")

    doc.add_heading("六、下周计划", level=1)
    add_bullet(doc, "针对 cuff、pocket 的低 IoU 样本做按区域 failure review，区分服饰缺失、遮挡、左右语义和小目标定位失败。")
    add_bullet(doc, "继续优化 semantic prompt 与 side-aware 视觉定位；对 zipper、decoration 等目标先做离线比较，不直接修改默认路径。")
    add_bullet(doc, "若后续尝试 DINOv2/CLIP 区域特征或更强 grounding 模型，仍以人工 benchmark 和 image-held-out 验证为准。")
    add_bullet(doc, "只有同时满足人工 IoU 改善和延迟可接受，才考虑把实验分支提升为可选线上 backend。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
