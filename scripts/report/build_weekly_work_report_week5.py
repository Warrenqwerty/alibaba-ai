from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("reports/服饰实例分割项目周报-第五周.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    paragraph = add_paragraph(doc, text, "List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(1)

    for style_name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(1)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("本周工作汇报：服饰细粒度视觉模块（3.1.2 方向复盘与重构）")
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(16)
    run.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("汇报周期：第五周｜方向：多模态电商服饰细粒度视觉基础模块")
    meta_run.font.name = "SimSun"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    meta_run.font.size = Pt(10.5)
    meta_run.font.color.rgb = RGBColor(90, 90, 90)


def add_feedback_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["导师反馈", "暴露的问题", "本周调整", "当前结论"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")

    rows = [
        (
            "训练/评估都来自 landmark pseudo-label + rule fallback",
            "伪标签本身有噪声，继续优化 pseudo-label IoU 可能跑偏",
            "建立人工 bbox benchmark，用人工框作为独立评估标准",
            "弱标签指标只保留为开发诊断，不再作为 PRD 结论",
        ),
        (
            "候选区域生成利用 DeepFashion2 GT mask 和 landmark 先验",
            "候选级验证结果偏乐观，存在数据泄露风险",
            "把 candidate-listwise ranker 从 online path 暂时关掉",
            "线上默认改为 heuristic-only，learned ranker 作为实验分支",
        ),
        (
            "完整 pipeline 脱离 landmark 后明显退化",
            "训练环境和真实推理环境不一致",
            "用 3.1.1 预测结果 + 人工标注重新评估",
            "manual benchmark 才是当前最可信的判断依据",
        ),
        (
            "文档提到 DINOv2/CLIP 图文匹配",
            "前期过度依赖 DeepFashion2 弱监督训练",
            "重构 3.1.2 计划，回到预训练 grounding / 图文匹配路线",
            "后续优先验证 GroundingDINO/OWL-ViT/CLIP 类 baseline",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.5)


def add_manual_benchmark_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["阶段", "评估规模", "avg bbox IoU", "Hit@0.3 / Hit@0.5", "结论"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")

    rows = [
        ("初始 heuristic", "55 labeled", "0.2544", "0.4000 / 0.2545", "优于 learned hybrid"),
        ("candidate-listwise hybrid", "55 labeled", "0.2324", "0.3455 / 0.2000", "人工评估未带来收益"),
        ("合并人工集 baseline", "122 labeled", "0.2812", "0.4344 / 0.2623", "作为后续比较基线"),
        ("failure review 后 refinement", "122 labeled", "0.3064", "0.4754 / 0.2787", "pocket/waist 明显提升"),
        ("cuff variant refinement", "122 labeled", "0.3123", "0.4836 / 0.2705", "整体继续提升，袖口仍是瓶颈"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)


def add_region_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["问题区域", "主要失败原因", "本周修正", "结果/判断"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")

    rows = [
        (
            "cuff",
            "短袖/袖窿和长袖袖口混在一起，纯几何框容易选到整条袖子",
            "增加 upper-sleeve 与 lower-terminal cuff 两类候选",
            "IoU 0.0190 -> 0.0904，但仍然低，说明需要视觉 grounding",
        ),
        (
            "pocket",
            "左右方向按图像坐标理解，和服饰/穿着者左右不一致",
            "side-specific pocket 改为 wearer/garment left-right convention",
            "IoU 0.0000 -> 0.1337，方向修正有效",
        ),
        (
            "waist",
            "不同类别腰部位置差异大，裤子/裙子/连衣裙不能用同一纵向窗口",
            "根据 garment category 使用不同 upper-band/waist-band 几何",
            "IoU 0.0961 -> 0.2306，类别感知规则有效",
        ),
        (
            "zipper / pattern / decoration",
            "目标小且依赖真实视觉纹理，规则候选不稳定",
            "暂不继续强行规则调参，放入 pretrained grounding 下一阶段验证",
            "更适合用 GroundingDINO/CLIP/SAM 类图文定位能力处理",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.5)


def add_direction_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["路线", "之前做法", "问题", "调整后做法"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")

    rows = [
        (
            "训练数据",
            "DeepFashion2 mask + landmark 构造 pseudo-label",
            "没有 query-level 人工真值，容易学到几何先验",
            "只作为弱监督/诊断；主评估转向人工 bbox benchmark",
        ),
        (
            "排序模型",
            "hash ranker / listwise context ranker",
            "候选级指标好，但 manual/full pipeline 不稳定",
            "从 online path 关掉，保留为 archived experiment",
        ),
        (
            "文档路线",
            "先尝试用 DeepFashion2 弱监督训练",
            "和 PRD 中 DINOv2/CLIP 图文匹配方向不完全一致",
            "下一步做 pretrained grounding baseline",
        ),
        (
            "线上策略",
            "可选传入 --ranker-checkpoint",
            "误传 checkpoint 可能让结果变差",
            "默认 heuristic-only，只有超过 manual benchmark 才接入新 backend",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.5)


def build_report() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("一、本周主要工作", level=1)
    add_paragraph(
        doc,
        "本周主要围绕 3.1.2「语言引导的局部区域定位」做方向复盘和评估体系重构。根据导师反馈，前几周的训练和评估过度依赖 "
        "landmark pseudo-label + rule fallback，缺少人工真实标注作为基准，容易把模型优化到伪标签几何上。因此本周重点从继续训练 "
        "weak ranker，调整为建立小规模人工 benchmark、复查 PRD 中 DINOv2/CLIP 的图文匹配路线，并把后续主线改为 pretrained grounding。"
    )

    doc.add_heading("二、导师反馈与本周调整", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["反馈/问题", "本周处理", "当前结论"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    rows = [
        (
            "训练/评估都来自伪标签，可能围绕 noisy pseudo-label 跑偏",
            "建立人工 bbox benchmark，不使用 landmark 标注",
            "manual benchmark 作为 3.1.2 主要评估依据",
        ),
        (
            "候选区域利用 GT mask 和 landmark 先验，候选级结果偏乐观",
            "把 candidate-listwise ranker 从 online path 暂时关掉",
            "线上默认 heuristic-only，learned ranker 只保留为实验分支",
        ),
        (
            "PRD 更接近 DINOv2/CLIP 区域特征与文本特征匹配路线",
            "重构 README、AutoDL setup 和 3.1.2 plan",
            "下一步优先验证 GroundingDINO/OWL-ViT/CLIP 类 baseline",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.5)

    doc.add_heading("三、Manual Benchmark 与策略修正", level=1)
    add_paragraph(
        doc,
        "本周完成两轮人工标注合并，共得到 122 条有效 labeled records。标注时只根据图片和 query 手动画 bbox，"
        "不可标注 query 标为 unlabeled，不再用 DeepFashion2 landmark 作为评估真值。"
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["阶段", "规模", "avg bbox IoU", "Hit@0.3 / Hit@0.5", "结论"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    rows = [
        ("初始 heuristic", "55", "0.2544", "0.4000 / 0.2545", "优于 learned hybrid"),
        ("candidate-listwise hybrid", "55", "0.2324", "0.3455 / 0.2000", "人工评估未带来收益"),
        ("合并人工集 baseline", "122", "0.2812", "0.4344 / 0.2623", "作为后续比较基线"),
        ("failure review + rule refinement", "122", "0.3123", "0.4836 / 0.2705", "整体提升，cuff 仍是瓶颈"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.5)
    add_paragraph(
        doc,
        "Failure review 后主要修正了三类问题：cuff 增加 upper-sleeve / lower-terminal 候选；pocket 左右方向改为服饰/穿着者左右；"
        "waist 根据 garment category 使用不同纵向区域。pocket 和 waist 提升明显，但 cuff 仍然较弱，说明纯几何规则接近上限。"
    )

    doc.add_heading("四、当前判断", level=1)
    add_bullet(doc, "DeepFashion2 有 mask、bbox、category 和 landmark，但没有 query-level 语言区域人工标注，不能单独支撑 3.1.2 的完整训练。")
    add_bullet(doc, "weak-label ranker 可以作为历史实验和辅助诊断，但不能作为 PRD 级语言 grounding 的主要结论。")
    add_bullet(doc, "当前 online policy 固定为 heuristic-only；只有新模型超过 manual benchmark，才考虑接入 online path。")
    add_bullet(doc, "3.1.2 后续主线应回到预训练 grounding / 图文匹配：GroundingDINO、OWL-ViT/OWL-V2、Chinese-CLIP 或 CLIP prompt mapping。")

    doc.add_heading("五、下周计划", level=1)
    add_bullet(doc, "实现 pretrained grounding 离线评估脚本，输入人工标注 JSONL，输出预测 bbox、IoU、可视化和 summary。")
    add_bullet(doc, "优先验证 GroundingDINO 或 OWL-ViT/OWL-V2；若模型偏英文，增加中文 query 到英文 prompt 的模板映射。")
    add_bullet(doc, "继续保留 heuristic-only 作为 control baseline；新模型必须超过 manual benchmark 才能进入 online path。")
    add_bullet(doc, "重点比较 cuff、pocket、zipper、pattern 等规则困难区域；如需继续标注，只做小规模、有目标的补充标注。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
